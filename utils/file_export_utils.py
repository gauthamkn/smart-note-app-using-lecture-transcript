from fpdf import FPDF
from docx import Document
import tempfile
import os

def generate_pdf(title, summary, notes, flashcards):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, txt=title, ln=True, align="C")

    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt="Summary", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 10, txt=summary)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt="Topic-wise Notes", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 10, txt=notes)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt="Flashcards", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 10, txt=flashcards)

    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_pdf.name)
    return temp_pdf.name


def generate_docx(title, summary, notes, flashcards):
    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)

    doc.add_heading("Topic-wise Notes", level=2)
    doc.add_paragraph(notes)

    doc.add_heading("Flashcards", level=2)
    doc.add_paragraph(flashcards)

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    return temp_docx.name
